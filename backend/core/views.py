from django.utils.timezone import now
from rest_framework import viewsets, status, permissions
from rest_framework.response import Response
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework_simplejwt.views import TokenObtainPairView
from .models import User, Classroom, Enrollment, Assignment, Submission, AIReport, Announcement, Notification
from .serializers import (
    UserSerializer, ClassroomSerializer, EnrollmentSerializer, 
    AssignmentSerializer, SubmissionSerializer, MyTokenObtainPairSerializer,
    AnnouncementSerializer, NotificationSerializer
)
import random
import string
import os
from pathlib import Path
from google import genai
import environ
from spellchecker import SpellChecker
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import spacy
import pdfplumber
import docx
import textstat
import io

# Initialize environ
env = environ.Env()
environ.Env.read_env(os.path.join(Path(__file__).resolve().parent.parent, '.env'))

# Configure AI API (Groq - Free alternative to Gemini)
GROQ_API_KEY = env('GROQ_API_KEY', default=None)
if GROQ_API_KEY:
    from groq import Groq
    groq_client = Groq(api_key=GROQ_API_KEY)
else:
    groq_client = None

# Configure Gemini API (new google.genai package)
GEMINI_API_KEY = env('GEMINI_API_KEY', default=None)
if GEMINI_API_KEY:
    genai_client = genai.Client(api_key=GEMINI_API_KEY)
else:
    genai_client = None

nlp = spacy.load("en_core_web_sm")

@api_view(['GET'])
@permission_classes([permissions.AllowAny])
def api_root(request):
    return Response({
        "message": "Welcome to AI Classroom API",
        "status": "Running",
        "endpoints": {
            "auth": "/api/token/",
            "classrooms": "/api/classrooms/",
            "users": "/api/users/"
        }
    })

class MyTokenObtainPairView(TokenObtainPairView):
    serializer_class = MyTokenObtainPairSerializer

class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer
    permission_classes = [permissions.AllowAny]

    @action(detail=False, methods=['get', 'put', 'patch'], permission_classes=[permissions.IsAuthenticated])
    def me(self, request):
        if request.method == 'GET':
            serializer = self.get_serializer(request.user)
            return Response(serializer.data)
        
        serializer = self.get_serializer(request.user, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data)

class ClassroomViewSet(viewsets.ModelViewSet):
    queryset = Classroom.objects.all()
    serializer_class = ClassroomSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        if user.role == 'faculty':
            return Classroom.objects.filter(faculty=user)
        # Students see classrooms they are enrolled in
        return Classroom.objects.filter(enrollments__student=user)

    def perform_create(self, serializer):
        # Generate unique class code
        while True:
            code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
            if not Classroom.objects.filter(class_code=code).exists():
                break
        serializer.save(faculty=self.request.user, class_code=code)

    @action(detail=False, methods=['post'])
    def join(self, request):
        code = request.data.get('code')
        try:
            classroom = Classroom.objects.get(class_code=code)
            Enrollment.objects.get_or_create(classroom=classroom, student=request.user)
            return Response(ClassroomSerializer(classroom).data)
        except Classroom.DoesNotExist:
            return Response({'error': 'Invalid class code'}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=['post'])
    def global_chat(self, request):
        try:
            user = request.user
            user_query = request.data.get('message', '')
            
            # Aggregate data from all user's classrooms
            if user.role == 'faculty':
                classrooms = Classroom.objects.filter(faculty=user)
            else:
                classrooms = Classroom.objects.filter(enrollments__student=user)
                
            assignments = Assignment.objects.filter(classroom__in=classrooms)
            announcements = Announcement.objects.filter(classroom__in=classrooms).order_by('-created_at')
            my_submissions = Submission.objects.filter(assignment__in=assignments, student=user)
            
            pending_count = max(0, assignments.count() - my_submissions.count())
            
            if groq_client:
                try:
                    response = groq_client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[
                            {"role": "system", "content": "You are a helpful AI assistant. Answer clearly using bullet points."},
                            {"role": "user", "content": user_query}
                        ],
                        temperature=0.7,
                        max_tokens=1000
                    )
                    response_text = response.choices[0].message.content
                    return Response({'response': response_text})
                except Exception as ai_err:
                    return Response({'response': f"AI Engine Error: {str(ai_err)}", 'debug': True}, status=status.HTTP_200_OK)
            
            # Fallback if no API key
            return Response({
                'response': f"Hello {user.username}! I am currently running in limited mode because the Gemini API key is not configured. "
                            f"You have {classrooms.count()} classrooms and {pending_count} pending assignments across all your classes."
            })
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            print(error_trace) # Log to server terminal
            return Response({
                'response': f"System Error: {str(e)}", 
                'trace': error_trace if os.environ.get('DEBUG') else None
            }, status=status.HTTP_400_BAD_REQUEST) # Use 400 instead of 500 to see message in frontend

    @action(detail=True, methods=['get'])
    def stream(self, request, pk=None):
        classroom = self.get_object()
        announcements = Announcement.objects.filter(classroom=classroom).order_by('-created_at')
        assignments = Assignment.objects.filter(classroom=classroom).order_by('-created_at')
        
        # Combine and sort? Or return separate? Let's return separate for now.
        return Response({
            'announcements': AnnouncementSerializer(announcements, many=True).data,
            'assignments': AssignmentSerializer(assignments, many=True).data
        })
    @action(detail=True, methods=['get'])
    def people(self, request, pk=None):
        classroom = self.get_object()
        enrollments = Enrollment.objects.filter(classroom=classroom).select_related('student')
        return Response({
            'faculty': UserSerializer(classroom.faculty).data,
            'students': UserSerializer([e.student for e in enrollments], many=True).data
        })

    @action(detail=True, methods=['get'])
    def grades(self, request, pk=None):
        classroom = self.get_object()
        assignments = Assignment.objects.filter(classroom=classroom)
        enrollments = Enrollment.objects.filter(classroom=classroom).select_related('student')
        
        gradebook = []
        for enrollment in enrollments:
            student = enrollment.student
            student_grades = []
            for assignment in assignments:
                submission = Submission.objects.filter(student=student, assignment=assignment).first()
                student_grades.append({
                    'assignment_id': assignment.id,
                    'grade': submission.grade if submission else None,
                    'status': submission.status if submission else 'missing'
                })
            gradebook.append({
                'student_id': student.id,
                'student_name': student.username,
                'grades': student_grades
            })
            
        return Response({
            'assignments': AssignmentSerializer(assignments, many=True).data,
            'gradebook': gradebook
        })

    @action(detail=True, methods=['get'])
    def analytics(self, request, pk=None):
        classroom = self.get_object()
        assignments = Assignment.objects.filter(classroom=classroom)
        enrollments_count = Enrollment.objects.filter(classroom=classroom).count()
        
        stats = []
        for asgn in assignments:
            subs = Submission.objects.filter(assignment=asgn)
            avg_grade = sum([s.grade for s in subs if s.grade]) / subs.filter(grade__isnull=False).count() if subs.filter(grade__isnull=False).exists() else 0
            stats.append({
                'assignment_id': asgn.id,
                'title': asgn.title,
                'submission_rate': (subs.count() / enrollments_count * 100) if enrollments_count > 0 else 0,
                'average_grade': round(avg_grade, 2)
            })
            
        return Response(stats)

    @action(detail=True, methods=['post'])
    def ai_chat(self, request, pk=None):
        try:
            classroom = self.get_object()
            user = request.user
            user_query = request.data.get('message', '')
            
            # Real-time data fetching for context
            assignments = Assignment.objects.filter(classroom=classroom)
            announcements = Announcement.objects.filter(classroom=classroom).order_by('-created_at')
            my_submissions = Submission.objects.filter(assignment__classroom=classroom, student=user)
            graded_submissions = my_submissions.filter(status='graded')
            
            # Prepare context for Gemini
            context = f"Classroom: {classroom.name} ({classroom.section})\n"
            context += f"Instructor: {classroom.faculty.username}\n"
            context += f"Current User: {user.username} (Role: {user.role})\n\n"
            
            context += "Available Assignments:\n"
            for a in assignments:
                status_str = "Submitted" if my_submissions.filter(assignment=a).exists() else "Pending"
                context += f"- {a.title}: {a.points} points, Due {a.due_date.strftime('%Y-%m-%d %H:%M')}. Status: {status_str}\n"
                
            context += "\nRecent Announcements:\n"
            for ann in announcements[:3]:
                context += f"- [{ann.created_at.strftime('%Y-%m-%d')}] {ann.content}\n"
                
            if graded_submissions.exists():
                grades = [s.grade for s in graded_submissions if s.grade is not None]
                if grades:
                    avg = sum(grades) / len(grades)
                    context += f"\nUser Performance: {len(grades)} graded tasks, Average Score: {avg:.1f} points.\n"

            if genai_client:
                try:
                    system_instruction = (
                        f"You are the Classroom AI Assistant for '{classroom.name}'. "
                        "You have direct access to assignments, student grades, and announcements. "
                        "If a student asks about their grade, refer to the performance data. "
                        "If they ask about deadlines, list the upcoming ones. "
                        "Always be encouraging, precise, and educational. Use markdown."
                    )
                    full_prompt = (
                        f"{system_instruction}\n\n"
                        f"### Context: {classroom.name}\n"
                        f"{context}\n\n"
                        f"### User Message\n{user_query}"
                    )
                    
                    response = genai_client.models.generate_content(
                        model='gemini-1.5-flash',
                        contents=full_prompt
                    )
                    response_text = response.text if response.text else "I failed to process that request. Try asking about specific classroom details like deadlines or your grades."
                    return Response({'response': response_text})
                except Exception as ai_err:
                    return Response({'response': f"Classroom AI Error: {str(ai_err)}"}, status=status.HTTP_200_OK)
            
            # Fallback to local logic if no API key
            message = user_query.lower()
            if 'assignment' in message or 'work' in message:
                pending = assignments.count() - my_submissions.count()
                titles = ", ".join([a.title for a in assignments[:3]])
                response = f"There are {assignments.count()} assignments in {classroom.name}. "
                if pending > 0:
                    response += f"You have {pending} pending tasks. "
                response += f"The latest ones include: {titles}."
            elif 'deadline' in message or 'due' in message:
                upcoming = assignments.filter(due_date__gte=now()).order_by('due_date').first()
                if upcoming:
                    response = f"The next deadline is for '{upcoming.title}' on {upcoming.due_date.strftime('%B %d at %H:%M')}."
                else:
                    response = "Good news! There are no upcoming deadlines at the moment."
            else:
                response = "I'm currently running in low-power mode. Please add a Gemini API Key to enable full AI intelligence!"
                
            return Response({'response': response})
        except Exception as e:
            return Response({'response': f"System Error: {str(e)}"}, status=status.HTTP_400_BAD_REQUEST)

class AssignmentViewSet(viewsets.ModelViewSet):
    queryset = Assignment.objects.all()
    serializer_class = AssignmentSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        classroom_id = self.request.query_params.get('classroom')
        if classroom_id:
            return Assignment.objects.filter(classroom_id=classroom_id)
        return Assignment.objects.all()

    def perform_create(self, serializer):
        # Ensure classroom belongs to the faculty
        classroom_id = self.request.data.get('classroom')
        assignment = serializer.save(classroom_id=classroom_id)
        
        # Notify all enrolled students
        enrollments = Enrollment.objects.filter(classroom_id=classroom_id)
        for enrollment in enrollments:
            Notification.objects.create(
                user=enrollment.student,
                title="New Assignment",
                message=f"A new assignment '{assignment.title}' has been posted in {assignment.classroom.name}.",
                link=f"/classroom/{classroom_id}/assignment/{assignment.id}"
            )

class SubmissionViewSet(viewsets.ModelViewSet):
    queryset = Submission.objects.all()
    serializer_class = SubmissionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        queryset = Submission.objects.all().select_related('student', 'assignment')
        assignment_id = self.request.query_params.get('assignment')
        
        if assignment_id:
            queryset = queryset.filter(assignment_id=assignment_id)
            
        if user.role == 'student':
            queryset = queryset.filter(student=user)
        # If faculty, they see all submissions for classrooms they lead
        elif user.role == 'faculty':
            queryset = queryset.filter(assignment__classroom__faculty=user)
            
        return queryset

    def perform_create(self, serializer):
        submission = serializer.save(student=self.request.user)
        self.analyze_ai(submission)
        
        # Notify Faculty
        Notification.objects.create(
            user=submission.assignment.classroom.faculty,
            title="New Submission",
            message=f"{submission.student.username} submitted '{submission.assignment.title}'.",
            link=f"/classroom/{submission.assignment.classroom.id}/assignment/{submission.assignment.id}"
        )

    def perform_update(self, serializer):
        old_status = self.get_object().status
        old_file = self.get_object().file
        submission = serializer.save()
        
        # Only trigger AI Analysis if the file was actually changed
        # Don't re-analyze when faculty is just grading
        if submission.file != old_file:
            self.analyze_ai(submission)

        # Notify student if graded
        if old_status != 'graded' and submission.status == 'graded':
            Notification.objects.create(
                user=submission.student,
                title="Assignment Graded",
                message=f"Your submission for '{submission.assignment.title}' has been graded: {submission.grade}/{submission.assignment.points}.",
                link=f"/classroom/{submission.assignment.classroom.id}/assignment/{submission.assignment.id}"
            )

    def analyze_ai(self, submission):
        try:
            content = ""
            metadata = {"type": "unknown", "pages": 0, "styles": []}
            file_path = submission.file.path
            _, ext = os.path.splitext(file_path)

            if ext.lower() == '.pdf':
                metadata["type"] = "pdf"
                with pdfplumber.open(file_path) as pdf:
                    metadata["pages"] = len(pdf.pages)
                    extracted_text_parts = []
                    style_samples = []
                    
                    for i, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text: extracted_text_parts.append(text)
                        
                        # Extract rudimentary style info from the first few chars of each page
                        # to give AI a hint about fonts/margins
                        if i < 3: # Check first 3 pages
                            words = page.extract_words(extra_attrs=['fontname', 'size', 'x0', 'top'])
                            if words:
                                style_samples.append({
                                    "page": i + 1,
                                    "sample_fonts": list(set([w['fontname'] for w in words[:10]])),
                                    "sample_sizes": list(set([w['size'] for w in words[:10]])),
                                    "margin_left_sample": min([w['x0'] for w in words]),
                                    "margin_top_sample": min([w['top'] for w in words])
                                })
                    content = "\n".join(extracted_text_parts)
                    metadata["styles"] = style_samples

            elif ext.lower() in ['.docx', '.doc']:
                metadata["type"] = "docx"
                doc = docx.Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
                
                # Extract DOCX style info
                styles_found = []
                for i, para in enumerate(doc.paragraphs[:20]): # Check first 20 paragraphs
                    para_style = {
                        "alignment": str(para.alignment) if para.alignment else "Left/Default",
                        "fonts": [],
                        "sizes": []
                    }
                    for run in para.runs:
                        if run.font.name: para_style["fonts"].append(run.font.name)
                        if run.font.size: para_style["sizes"].append(str(run.font.size.pt))
                    
                    if para.text.strip():
                        styles_found.append(para_style)
                metadata["styles"] = styles_found

            else:
                submission.file.open(mode='rb')
                content = submission.file.read().decode('utf-8', errors='ignore')
                submission.file.close()

            if not content or not content.strip():
                content = "Empty or unreadable file content."

            # Spelling (with location tracking and proper noun filtering)
            spell = SpellChecker()
            import re
            
            # Common British/US spelling variants to accept
            british_variants = {
                'analyse', 'behaviour', 'colour', 'favour', 'honour', 'labour',
                'neighbour', 'rumour', 'splendour', 'centre', 'metre', 'litre',
                'defence', 'licence', 'offence', 'pretence', 'organisation',
                'realise', 'recognise', 'specialise', 'summarise', 'optimise'
            }
            skip_words = {'aiml', 'proteus', 'matlab', 'github', 'api', 'json', 'html', 'css', 'dsp', 'fsk'}
            
            # Use spaCy to detect proper nouns (names, organizations, locations)
            doc_nlp = nlp(content[:5000])  # Process first 5000 chars for NER
            proper_nouns = set()
            for ent in doc_nlp.ents:
                # Skip entities that are names, organizations, or locations
                if ent.label_ in ['PERSON', 'ORG', 'GPE', 'LOC', 'FAC']:
                    # Add each word in the entity to proper nouns set
                    for token in ent.text.split():
                        cleaned_token = re.sub(r'[^a-zA-Z]', '', token).lower()
                        if cleaned_token:
                            proper_nouns.add(cleaned_token)
            
            # Split content into lines for location tracking
            lines = content.split('\n')
            misspelled = []
            detailed_errors = []
            
            for line_num, line in enumerate(lines[:100], start=1):  # Check first 100 lines
                words_in_line = line.split()
                
                for word in words_in_line:
                    # Clean the word
                    cleaned = re.sub(r'[^a-zA-Z]', '', word).lower()
                    
                    # Only check words 4+ characters
                    if len(cleaned) >= 4 and cleaned.isalpha():
                        # Check if misspelled
                        if cleaned in spell.unknown([cleaned]):
                            # Filter out false positives
                            if (cleaned not in skip_words and 
                                cleaned not in british_variants and 
                                cleaned not in proper_nouns):  # NEW: Skip proper nouns
                                
                                # Get suggestion
                                correction = spell.correction(cleaned)
                                
                                # Add to simple list (for backward compatibility)
                                if cleaned not in misspelled:
                                    misspelled.append(cleaned)
                                
                                # Add detailed error info
                                detailed_errors.append({
                                    'word': cleaned,
                                    'line': line_num,
                                    'context': line.strip()[:100],  # First 100 chars of line
                                    'suggestion': correction if correction != cleaned else None
                                })
                
                # Limit to 20 detailed errors to avoid overwhelming the student
                if len(detailed_errors) >= 20:
                    break

            # Grammar/Clarity
            doc = nlp(content[:3000])
            suggestions = []
            grammar_errors = []
            
            for sent in doc.sents:
                if len(sent.text.split()) > 25:
                    suggestions.append(f"Long sentence: {sent.text[:40]}...")
                if any(tok.dep_ == 'nsubjpass' for tok in sent):
                    grammar_errors.append(f"Passive voice: {sent.text[:40]}...")

            # Readability (Safe wrap)
            try:
                if content.strip():
                    import nltk
                    try:
                        nltk.data.find('corpora/cmudict')
                    except LookupError:
                        nltk.download('cmudict', quiet=True)
                    readability = textstat.flesch_reading_ease(content[:5000])
                else:
                    readability = 0
            except Exception:
                # If textstat fails, default to 0
                readability = 0
            
            # Plagiarism (Optimized for limited RAM)
            similarity_score = 0.0
            matched_text = ""
            
            # Only run plagiarism check if there aren't too many submissions
            # and limit the comparison to first 10 to prevent OOM
            other_submissions = Submission.objects.filter(assignment=submission.assignment).exclude(id=submission.id)
            
            if other_submissions.exists() and len(content) > 100:
                texts = [content[:5000]] # Compare first 5000 chars
                for other in other_submissions[:10]: 
                    try:
                        # For simplicity, we only compare against text-based submissions 
                        # or skip deep parsing here to save memory. 
                        # Ideally, extracted text should be stored in the DB.
                        if other.file.name.lower().endswith('.txt'):
                            other.file.open('rb')
                            texts.append(other.file.read().decode('utf-8', errors='ignore')[:5000])
                            other.file.close()
                    except:
                        continue
                
                if len(texts) > 1:
                    try:
                        vectorizer = TfidfVectorizer(stop_words='english', max_features=1000).fit_transform(texts)
                        vectors = vectorizer.toarray()
                        cosine_sim = cosine_similarity(vectors)
                        max_sim = max(cosine_sim[0][1:]) if len(cosine_sim[0]) > 1 else 0
                        similarity_score = round(max_sim * 100, 2)
                        matched_text = "Similar patterns found." if similarity_score > 40 else ""
                    except Exception as plag_err:
                        print(f"Plagiarism check error: {plag_err}")

            # Formatting & Instruction Check (Gemini)
            formatting_result = {}


            # Formatting Check (AI-powered with fallback)
            if submission.assignment.formatting_instructions:
                instructions_lower = submission.assignment.formatting_instructions.lower()
                
                # Basic rule-based check (works without AI)
                basic_issues = []
                basic_score = 100
                
                # Check for common formatting requirements
                if 'times new roman' in instructions_lower:
                    fonts_found = metadata.get('styles', [{}])[0].get('fonts', []) if metadata.get('styles') else []
                    if fonts_found and not any('times' in str(f).lower() for f in fonts_found):
                        basic_issues.append("Font may not be Times New Roman")
                        basic_score -= 20
                
                if '12' in instructions_lower and ('font' in instructions_lower or 'size' in instructions_lower):
                    sizes_found = metadata.get('styles', [{}])[0].get('sizes', []) if metadata.get('styles') else []
                    if sizes_found and not any('12' in str(s) for s in sizes_found):
                        basic_issues.append("Font size may not be 12pt")
                        basic_score -= 20
                
                # Try AI enhancement using Groq (free & fast)
                if groq_client:
                    try:
                        system_instruction = (
                            "You are an academic formatting assistant. Check if the document follows the teacher's instructions. "
                            "Return ONLY a valid JSON object with keys: 'compliant' (boolean), 'issues' (array of strings), 'score' (number 0-100). "
                            "Be strict but fair."
                        )
                        
                        user_prompt = (
                            f"Teacher Instructions: {submission.assignment.formatting_instructions}\n\n"
                            f"Document Metadata: {metadata}\n\n"
                            f"Text Sample: {content[:500]}\n\n"
                            "Analyze and return JSON only."
                        )
                        
                        response = groq_client.chat.completions.create(
                            model="llama-3.3-70b-versatile",  # Fast, free Llama model
                            messages=[
                                {"role": "system", "content": system_instruction},
                                {"role": "user", "content": user_prompt}
                            ],
                            temperature=0.3,
                            max_tokens=500
                        )
                        
                        json_str = response.choices[0].message.content.strip()
                        json_str = json_str.replace('```json', '').replace('```', '').strip()
                        import json
                        formatting_result = json.loads(json_str)
                        
                    except Exception as e:
                        # AI failed - use basic check results
                        formatting_result = {
                            "compliant": basic_score >= 70,
                            "score": basic_score,
                            "issues": basic_issues if basic_issues else ["Basic formatting check completed (AI unavailable)"]
                        }
                else:
                    # No API key - use basic check
                    formatting_result = {
                        "compliant": basic_score >= 70,
                        "score": basic_score,
                        "issues": basic_issues if basic_issues else ["Basic formatting check completed"]
                    }

            # Feedback and Grade
            feedback_parts = []
            if similarity_score > 30: feedback_parts.append("High similarity.")
            if len(misspelled) > 5: feedback_parts.append("Spelling issues.")
            
            # Incorporate formatting into summary
            format_score = 100
            if formatting_result:
                if not formatting_result.get('compliant', True):
                    feedback_parts.append("Formatting needs improvement.")
                    format_score = formatting_result.get('score', 100)
            
            summary = " ".join(feedback_parts) if feedback_parts else "Good work."

            total_points = submission.assignment.points
            
            # Weighted scoring: 40% Logic/Grammar, 30% Plagiarism, 30% Formatting
            # Base deductions
            deduction_grammar = min(len(misspelled) * 2, 20) + min(len(grammar_errors) * 3, 15)
            deduction_plagiarism = 30 if similarity_score > 40 else 0
            
            # Calculate final components
            score_content = max(0, 100 - deduction_grammar)
            score_plagiarism = max(0, 100 - deduction_plagiarism)
            
            final_score_percent = (score_content * 0.4) + (score_plagiarism * 0.3) + (format_score * 0.3)
            suggested = max(0, int((final_score_percent / 100) * total_points))

            AIReport.objects.update_or_create(
                submission=submission,
                defaults={
                    'spelling_errors': misspelled,
                    'detailed_spelling_errors': detailed_errors,
                    'grammar_errors': grammar_errors,
                    'clarity_suggestions': suggestions,
                    'similarity_score': similarity_score,
                    'matched_text': matched_text,
                    'readability_score': readability,
                    'feedback_summary': summary,
                    'suggested_grade': suggested,
                    'formatting_analysis': formatting_result
                }
            )
        except Exception as e:
            # Fallback for errors so submission isn't blocked
            print(f"Analysis Error: {e}")
            AIReport.objects.get_or_create(
                submission=submission,
                defaults={
                    'feedback_summary': f"Analysis partially failed: {str(e)}",
                    'similarity_score': 0,
                    'detailed_spelling_errors': [],
                    'formatting_analysis': {"error": str(e)}
                }
            )

class AnnouncementViewSet(viewsets.ModelViewSet):
    queryset = Announcement.objects.all()
    serializer_class = AnnouncementSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        classroom_id = self.request.data.get('classroom')
        announcement = serializer.save(
            author=self.request.user,
            classroom_id=classroom_id
        )
        
        # Notify all enrolled students
        enrollments = Enrollment.objects.filter(classroom_id=classroom_id)
        for enrollment in enrollments:
            if enrollment.student != self.request.user:
                Notification.objects.create(
                    user=enrollment.student,
                    title="New Announcement",
                    message=f"{self.request.user.username} posted a new announcement in {announcement.classroom.name}.",
                    link=f"/classroom/{classroom_id}"
                )

class NotificationViewSet(viewsets.ModelViewSet):
    queryset = Notification.objects.all()
    serializer_class = NotificationSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        print(f"DEBUG: Notification fetch for user: {self.request.user}")
        if not self.request.user.is_authenticated:
            print("DEBUG: User not authenticated!")
        return Notification.objects.filter(user=self.request.user).order_by('-created_at')

    @action(detail=False, methods=['post'])
    def mark_all_as_read(self, request):
        Notification.objects.filter(user=request.user, is_read=False).update(is_read=True)
        return Response({'status': 'ok'})

# Import error correction module
from .error_correction import FileCorrector
from django.http import FileResponse

@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def detect_and_correct_errors(request):
    """
    Detect errors in uploaded file and return corrected version
    """
    file = request.FILES.get('file')
    
    if not file:
        return Response({
            'error': 'No file uploaded'
        }, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        corrector = FileCorrector()
        result, error = corrector.process_file(file)
        
        if error:
            return Response({
                'error': error
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Store the corrected file temporarily for download
        corrected_file = result['corrected_file']
        file_extension = result['file_extension']
        
        # Return error analysis and download info
        return Response({
            'success': True,
            'original_preview': result['original_text'],
            'corrected_preview': result['corrected_text'],
            'error_summary': result['error_summary'],
            'errors': {
                'spelling': result['errors']['spelling'][:10],  # Limit to 10 for preview
                'grammar': result['errors']['grammar'][:10],
                'formatting': result['errors']['formatting'][:10],
                'clarity': result['errors']['clarity'][:10]
            },
            'download_ready': True,
            'file_extension': file_extension
        })
    
    except Exception as e:
        return Response({
            'error': f'Error processing file: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def download_corrected_file(request):
    """
    Download the corrected file
    """
    file = request.FILES.get('file')
    
    if not file:
        return Response({
            'error': 'No file uploaded'
        }, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        corrector = FileCorrector()
        result, error = corrector.process_file(file)
        
        if error:
            return Response({
                'error': error
            }, status=status.HTTP_400_BAD_REQUEST)
        
        corrected_file = result['corrected_file']
        file_extension = result['file_extension']
        
        # Get original filename without extension
        original_name = os.path.splitext(file.name)[0]
        
        # Create response with file
        response = FileResponse(
            corrected_file,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_extension == 'docx' else 'text/plain'
        )
        response['Content-Disposition'] = f'attachment; filename="{original_name}_corrected.{file_extension}"'
        
        return response
    
    except Exception as e:
        return Response({
            'error': f'Error generating corrected file: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
