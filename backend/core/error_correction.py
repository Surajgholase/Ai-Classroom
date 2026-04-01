"""
AI-Powered Error Detection and File Correction System
Detects predefined errors and rewrites files with corrections
"""

import os
import re
import environ
from pathlib import Path
from io import BytesIO
# Heavy imports moved inside functions to optimize boot time

# Initialize environ
env = environ.Env()
environ.Env.read_env(os.path.join(Path(__file__).resolve().parent.parent, '.env'))

# Configure AI API
GROQ_API_KEY = env('GROQ_API_KEY', default=None)

if GROQ_API_KEY:
    from groq import Groq
    groq_client = Groq(api_key=GROQ_API_KEY)
else:
    groq_client = None


from .nlp_utils import get_nlp



class ErrorDetector:
    """Detects various types of errors in text"""
    
    def __init__(self):
        from spellchecker import SpellChecker
        self.spell_checker = SpellChecker()
        self._nlp = None

    @property
    def nlp(self):
        if self._nlp is None:
            from .nlp_utils import get_nlp
            self._nlp = get_nlp()
        return self._nlp
    
    def detect_spelling_errors(self, text):
        """Detect spelling errors with context"""
        words = text.split()
        errors = []
        seen_words = set()  # Track unique errors only
        
        for i, word in enumerate(words):
            # Clean word of punctuation for checking
            clean_word = re.sub(r'[^\w\s]', '', word.lower())
            
            # Skip very short words, numbers, and already seen words
            if not clean_word or len(clean_word) < 3 or clean_word.isdigit():
                continue
                
            if clean_word in seen_words:
                continue
            
            # Check if word is misspelled
            if clean_word not in self.spell_checker:
                correction = self.spell_checker.correction(clean_word)
                
                # Only add if we have a valid correction and it's different from original
                if correction and correction != clean_word:
                    errors.append({
                        'type': 'spelling',
                        'word': clean_word,
                        'position': i,
                        'correction': correction,
                        'context': ' '.join(words[max(0, i-3):min(len(words), i+4)])
                    })
                    seen_words.add(clean_word)
        
        return errors
    
    def detect_grammar_errors(self, text):
        """Detect basic grammar issues using spaCy"""
        doc = self.nlp(text)
        errors = []
        seen_issues = set()  # Track unique issues
        
        for sent in doc.sents:
            # Check for sentence fragments (very short sentences without verb)
            has_verb = any(token.pos_ == 'VERB' for token in sent)
            if len(sent) < 5 and not has_verb:
                issue_key = f"fragment_{sent.text[:20]}"
                if issue_key not in seen_issues:
                    errors.append({
                        'type': 'grammar',
                        'issue': 'sentence_fragment',
                        'text': sent.text,
                        'suggestion': 'Consider expanding this sentence or combining with another.'
                    })
                    seen_issues.add(issue_key)
            
            # Check for run-on sentences (very long sentences)
            if len(sent) > 40:
                issue_key = f"runon_{sent.text[:20]}"
                if issue_key not in seen_issues:
                    errors.append({
                        'type': 'grammar',
                        'issue': 'run_on_sentence',
                        'text': sent.text[:100] + '...' if len(sent.text) > 100 else sent.text,
                        'suggestion': 'Consider breaking this into multiple sentences.'
                    })
                    seen_issues.add(issue_key)
        
        # Check for repeated words in the entire text
        words = text.split()
        for i in range(len(words) - 1):
            word1 = re.sub(r'[^\w\s]', '', words[i].lower())
            word2 = re.sub(r'[^\w\s]', '', words[i + 1].lower())
            
            if word1 and word2 and word1 == word2 and len(word1) > 2:
                issue_key = f"repeated_{word1}_{i}"
                if issue_key not in seen_issues:
                    errors.append({
                        'type': 'grammar',
                        'issue': 'repeated_word',
                        'text': f"{words[i]} {words[i+1]}",
                        'suggestion': f'Remove duplicate word: "{word1}"'
                    })
                    seen_issues.add(issue_key)
        
        return errors
    
    def detect_formatting_errors(self, text):
        """Detect formatting issues"""
        errors = []
        seen_issues = set()
        
        # Check for multiple spaces (2 or more consecutive spaces)
        import re
        multiple_spaces = re.findall(r' {2,}', text)
        if multiple_spaces:
            # Count total occurrences
            total_count = len(multiple_spaces)
            issue_key = "multiple_spaces"
            if issue_key not in seen_issues:
                errors.append({
                    'type': 'formatting',
                    'issue': 'multiple_spaces',
                    'suggestion': f'Found {total_count} instance(s) of multiple spaces. Remove extra spaces between words.'
                })
                seen_issues.add(issue_key)
        
        # Check for missing punctuation at end of sentences
        lines = text.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            if line and len(line) > 10:
                # Check if line doesn't end with punctuation
                if not line[-1] in '.!?':
                    issue_key = f"missing_punct_{i}"
                    if issue_key not in seen_issues:
                        errors.append({
                            'type': 'formatting',
                            'issue': 'missing_punctuation',
                            'text': line[:50] + '...' if len(line) > 50 else line,
                            'suggestion': 'Add appropriate punctuation at the end.'
                        })
                        seen_issues.add(issue_key)
        
        # Check for inconsistent capitalization (sentences starting with lowercase)
        sentences = re.split(r'[.!?]\s+', text)
        for i, sent in enumerate(sentences):
            sent = sent.strip()
            if sent and len(sent) > 3:
                # Check if sentence starts with lowercase letter
                if sent[0].islower():
                    issue_key = f"capitalization_{i}"
                    if issue_key not in seen_issues:
                        errors.append({
                            'type': 'formatting',
                            'issue': 'capitalization',
                            'text': sent[:50] + '...' if len(sent) > 50 else sent,
                            'suggestion': 'Start sentence with capital letter.'
                        })
                        seen_issues.add(issue_key)
        
        return errors
    
    def detect_clarity_issues(self, text):
        """Detect clarity and readability issues"""
        import textstat
        errors = []
        
        # Check readability score
        try:
            reading_ease = textstat.flesch_reading_ease(text)
            if reading_ease < 30:
                errors.append({
                    'type': 'clarity',
                    'issue': 'low_readability',
                    'score': reading_ease,
                    'suggestion': 'Text is very difficult to read. Consider simplifying sentences.'
                })
        except:
            pass
        
        # Check for passive voice (simplified detection)
        doc = self.nlp(text)
        passive_count = 0
        for sent in doc.sents:
            for token in sent:
                if token.dep_ == 'auxpass':
                    passive_count += 1
                    errors.append({
                        'type': 'clarity',
                        'issue': 'passive_voice',
                        'text': sent.text,
                        'suggestion': 'Consider using active voice for clarity.'
                    })
        
        return errors
    
    def detect_all_errors(self, text):
        """Detect all types of errors"""
        all_errors = {
            'spelling': self.detect_spelling_errors(text),
            'grammar': self.detect_grammar_errors(text),
            'formatting': self.detect_formatting_errors(text),
            'clarity': self.detect_clarity_issues(text)
        }
        
        # Calculate error counts
        error_summary = {
            'total_errors': sum(len(errors) for errors in all_errors.values()),
            'spelling_count': len(all_errors['spelling']),
            'grammar_count': len(all_errors['grammar']),
            'formatting_count': len(all_errors['formatting']),
            'clarity_count': len(all_errors['clarity'])
        }
        
        return all_errors, error_summary


class FileCorrector:
    """Corrects errors in files and generates new versions"""
    
    def __init__(self):
        self.error_detector = ErrorDetector()
    
    def extract_text_from_file(self, file):
        """Extract text from uploaded file"""
        import pdfplumber
        import docx
        try:
            file_name = file.name.lower()
            
            if file_name.endswith('.pdf'):
                with pdfplumber.open(file) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    return '\n'.join(text_parts)
            
            elif file_name.endswith('.docx') or file_name.endswith('.doc'):
                doc = docx.Document(file)
                return '\n'.join([para.text for para in doc.paragraphs])
            
            elif file_name.endswith('.txt'):
                # Read the file content
                content = file.read().decode('utf-8', errors='ignore')
                # Reset file pointer for potential re-reading
                file.seek(0)
                return content
            
            else:
                return None, "Unsupported file format"
        
        except Exception as e:
            return None, f"Error reading file: {str(e)}"
    
    def correct_text_with_ai(self, text, errors):
        """Use AI to correct text based on detected errors"""
        if not groq_client:
            return self.correct_text_basic(text, errors)
        
        # Prepare error summary for AI
        error_summary = f"""
Detected Errors:
- Spelling errors: {len(errors.get('spelling', []))}
- Grammar errors: {len(errors.get('grammar', []))}
- Formatting errors: {len(errors.get('formatting', []))}
- Clarity issues: {len(errors.get('clarity', []))}

Specific Issues:
"""
        
        for error_type, error_list in errors.items():
            for error in error_list[:5]:  # Limit to first 5 of each type
                if error_type == 'spelling':
                    error_summary += f"\n- Spelling: '{error['word']}' → '{error['correction']}'"
                elif error_type == 'grammar':
                    error_summary += f"\n- Grammar: {error['issue']} - {error.get('suggestion', '')}"
                elif error_type == 'formatting':
                    error_summary += f"\n- Formatting: {error['issue']}"
                elif error_type == 'clarity':
                    error_summary += f"\n- Clarity: {error['issue']}"
        
        prompt = f"""You are an expert editor. Please correct the following text by fixing all errors while maintaining the original meaning and style.

{error_summary}

Original Text:
{text[:3000]}  

Instructions:
1. Fix all spelling errors
2. Correct grammar issues
3. Improve formatting and punctuation
4. Enhance clarity while keeping the original meaning
5. Maintain the original structure and tone
6. Return ONLY the corrected text, no explanations

Corrected Text:"""
        
        try:
            if groq_client:
                response = groq_client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[
                        {"role": "system", "content": "You are an expert editor who corrects text while maintaining its original meaning and style."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    max_tokens=4000
                )
                return response.choices[0].message.content.strip()
            
        
        except Exception as e:
            print(f"AI correction failed: {e}")
            return self.correct_text_basic(text, errors)
    
    def correct_text_basic(self, text, errors):
        """Basic text correction without AI"""
        corrected = text
        
        # Fix spelling errors
        for error in errors.get('spelling', []):
            if error['correction'] and error['correction'] != error['word']:
                # Use word boundary to avoid partial replacements
                pattern = r'\b' + re.escape(error['word']) + r'\b'
                corrected = re.sub(pattern, error['correction'], corrected, flags=re.IGNORECASE, count=1)
        
        # Fix repeated words
        for error in errors.get('grammar', []):
            if error.get('issue') == 'repeated_word':
                # Extract the repeated word
                text_parts = error.get('text', '').split()
                if len(text_parts) >= 2 and text_parts[0].lower() == text_parts[1].lower():
                    # Replace "word word" with "word"
                    pattern = r'\b' + re.escape(text_parts[0]) + r'\s+' + re.escape(text_parts[1]) + r'\b'
                    corrected = re.sub(pattern, text_parts[0], corrected, flags=re.IGNORECASE, count=1)
        
        # Fix formatting issues - multiple spaces
        corrected = re.sub(r'  +', ' ', corrected)  # Remove multiple spaces
        
        # Fix capitalization at start of lines
        lines = corrected.split('\n')
        corrected_lines = []
        for line in lines:
            line_stripped = line.strip()
            if line_stripped and line_stripped[0].islower():
                # Capitalize first letter
                line = line.replace(line_stripped[0], line_stripped[0].upper(), 1)
            corrected_lines.append(line)
        corrected = '\n'.join(corrected_lines)
        
        # Add missing punctuation at end of lines (if they look like sentences)
        lines = corrected.split('\n')
        corrected_lines = []
        for line in lines:
            line_stripped = line.strip()
            if line_stripped and len(line_stripped) > 10:
                # Check if line doesn't end with punctuation
                if not line_stripped[-1] in '.!?,;:':
                    line = line.rstrip() + '.'
            corrected_lines.append(line)
        corrected = '\n'.join(corrected_lines)
        
        return corrected
    
    def create_corrected_docx(self, original_text, corrected_text, errors, error_summary):
        """Create a new DOCX file with corrections"""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        
        # Add title
        title = doc.add_heading('AI-Corrected Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add error summary
        doc.add_heading('Error Summary', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Total Errors Found: {error_summary['total_errors']}\n").bold = True
        summary_para.add_run(f"• Spelling Errors: {error_summary['spelling_count']}\n")
        summary_para.add_run(f"• Grammar Errors: {error_summary['grammar_count']}\n")
        summary_para.add_run(f"• Formatting Errors: {error_summary['formatting_count']}\n")
        summary_para.add_run(f"• Clarity Issues: {error_summary['clarity_count']}\n")
        
        # Add corrected content
        doc.add_heading('Corrected Content', level=1)
        
        # Split into paragraphs and add to document
        paragraphs = corrected_text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.style.font.size = Pt(12)
        
        # Add detailed error list
        doc.add_page_break()
        doc.add_heading('Detailed Error List', level=1)
        
        for error_type, error_list in errors.items():
            if error_list:
                doc.add_heading(f'{error_type.capitalize()} Errors', level=2)
                for i, error in enumerate(error_list[:20], 1):  # Limit to 20 per type
                    error_para = doc.add_paragraph(style='List Number')
                    if error_type == 'spelling':
                        error_para.add_run(f"'{error['word']}' → '{error['correction']}'")
                    else:
                        error_para.add_run(f"{error.get('issue', 'Issue')}: {error.get('suggestion', error.get('text', ''))}")
        
        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
    
    def create_corrected_txt(self, corrected_text, errors, error_summary):
        """Create a corrected TXT file"""
        content = "=" * 60 + "\n"
        content += "AI-CORRECTED DOCUMENT\n"
        content += "=" * 60 + "\n\n"
        
        content += "ERROR SUMMARY\n"
        content += "-" * 60 + "\n"
        content += f"Total Errors Found: {error_summary['total_errors']}\n"
        content += f"• Spelling Errors: {error_summary['spelling_count']}\n"
        content += f"• Grammar Errors: {error_summary['grammar_count']}\n"
        content += f"• Formatting Errors: {error_summary['formatting_count']}\n"
        content += f"• Clarity Issues: {error_summary['clarity_count']}\n\n"
        
        content += "=" * 60 + "\n"
        content += "CORRECTED CONTENT\n"
        content += "=" * 60 + "\n\n"
        content += corrected_text + "\n\n"
        
        content += "=" * 60 + "\n"
        content += "DETAILED ERROR LIST\n"
        content += "=" * 60 + "\n\n"
        
        for error_type, error_list in errors.items():
            if error_list:
                content += f"\n{error_type.upper()} ERRORS:\n"
                content += "-" * 60 + "\n"
                for i, error in enumerate(error_list[:20], 1):
                    if error_type == 'spelling':
                        content += f"{i}. '{error['word']}' → '{error['correction']}'\n"
                    else:
                        content += f"{i}. {error.get('issue', 'Issue')}: {error.get('suggestion', error.get('text', ''))}\n"
        
        return BytesIO(content.encode('utf-8'))
    
    def process_file(self, file):
        """Main processing function"""
        # Extract text
        text = self.extract_text_from_file(file)
        if text is None:
            return None, "Failed to extract text from file"
        
        # Detect errors
        errors, error_summary = self.error_detector.detect_all_errors(text)
        
        # Correct text
        corrected_text = self.correct_text_with_ai(text, errors)
        
        # Create corrected file
        file_name = file.name.lower()
        if file_name.endswith('.docx') or file_name.endswith('.doc'):
            corrected_file = self.create_corrected_docx(text, corrected_text, errors, error_summary)
            file_extension = 'docx'
        else:
            corrected_file = self.create_corrected_txt(corrected_text, errors, error_summary)
            file_extension = 'txt'
        
        return {
            'original_text': text[:1000],  # First 1000 chars for preview
            'corrected_text': corrected_text[:1000],
            'errors': errors,
            'error_summary': error_summary,
            'corrected_file': corrected_file,
            'file_extension': file_extension
        }, None
